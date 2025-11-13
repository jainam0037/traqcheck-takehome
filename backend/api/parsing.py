"""
parsing.py
End-to-end resume parsing helpers:
- extract_text(): PDF/DOCX to plain text
- deterministic_extract(): regex/heuristics (email/phone/skills/role/company)
- llm_extract(): optional JSON from LLM
- merge_results(): combine rule-based + LLM with confidences
- update_candidate(): write extracted fields back to Candidate
"""

from __future__ import annotations
import os
import re
from typing import Dict, Tuple, List

from django.conf import settings

from .utils_text import normalize_space, truncate, canonical_phone, clamp01
from .schemas import Extracted, Confidence
from .llm_client import generate_structured

# -----------------------------------------------------------------------------
# Text extraction
# -----------------------------------------------------------------------------

def _extract_pdf(path: str) -> str:
    """Try PyPDF (pypdf) first, then fall back to pdfminer.six."""
    text = ""
    try:
        import pypdf  # type: ignore
        with open(path, "rb") as f:
            pdf = pypdf.PdfReader(f)
            pages = [page.extract_text() or "" for page in pdf.pages]
        text = "\n".join(pages)
    except Exception:
        text = ""

    if text and len(text.strip()) >= 50:
        return text

    try:
        from pdfminer.high_level import extract_text as pm_extract  # type: ignore
        text2 = pm_extract(path) or ""
        if len(text2.strip()) > len(text.strip()):
            return text2
    except Exception:
        pass

    return text or ""

def _extract_docx(path: str) -> str:
    """Extract text from a DOCX file, including table cells."""
    try:
        import docx  # python-docx
        d = docx.Document(path)
        lines = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        lines.append(cell.text)
        return "\n".join(lines)
    except Exception:
        return ""

def _normalize_keep_newlines(raw: str) -> str:
    """Preserve line breaks; normalize spaces inside each line."""
    raw = raw.replace("\r", "\n").replace("\x00", " ")
    lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in raw.split("\n")]
    lines = [ln for ln in lines if ln]
    return "\n".join(lines)

def extract_text(path: str, max_chars: int = 50000) -> str:
    """Extract textual content from a PDF or DOCX and truncate it."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        raw = _extract_pdf(path)
    elif ext == ".docx":
        raw = _extract_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    txt = _normalize_keep_newlines(raw)
    return truncate(txt, max_chars)

# -----------------------------------------------------------------------------
# Deterministic regex / heuristics
# -----------------------------------------------------------------------------

_EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")
# Flexible phone; we canonicalize after (7..15 digits)
_PHONE_RE = re.compile(r"(\+?\d[\d\s().\-]{6,}\d)")

_SKILL_KEYWORDS = {
    "python", "java", "javascript", "typescript", "react", "redux", "next.js", "node", "node.js",
    "django", "flask", "spring", "spring boot", "postgres", "postgresql", "mysql", "mssql",
    "mongodb", "neo4j", "redis", "kafka", "spark", "hadoop", "elasticsearch",
    "docker", "kubernetes", "eks", "aws", "gcp", "azure", "sagemaker",
    "celery", "graphql", "rest", "grpc", "pandas", "pytorch", "tensorflow", "scikit-learn",
    "selenium", "playwright", "jenkins", "github actions", "terraform", "prometheus", "grafana"
}

_MONTHS = r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)"
_SEP_CHARS = r"\|/•\-–—"

def _strip_trailing_meta(s: str) -> str:
    """Remove date ranges and trailing location/timeline separators from company."""
    s = re.split(rf"\s[{_SEP_CHARS}]\s", s, maxsplit=1)[0]
    s = re.sub(rf"\b{_MONTHS}\b\s+\d{{4}}\s*[–—\-]\s*\b{_MONTHS}\b\s+\d{{4}}", "", s, flags=re.I)
    s = re.sub(r"\b(19|20)\d{2}\s*[–—\-]\s*(19|20)\d{2}\b", "", s)  # 2021-2023
    s = re.sub(rf"\b{_MONTHS}\b\s+\d{{4}}", "", s, flags=re.I)       # Jun 2024
    return re.sub(r"\s{2,}", " ", s).strip()

def deterministic_extract(text: str) -> Dict[str, object]:
    """
    Pull low-hanging fruit using regex and simple heuristics.
    Returns a dict compatible with Extracted.model_dump().
    """
    out: Dict[str, object] = {"name": "", "email": "", "phone": "", "company": "", "designation": "", "skills": []}
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    # Email
    m = _EMAIL_RE.search(text)
    if m:
        out["email"] = m.group(0).lower()

    # Phone (prefer the first plausible)
    for m in _PHONE_RE.finditer(text):
        cand = canonical_phone(m.group(0))
        if cand:
            out["phone"] = cand
            break

    # Name heuristic: first non-meta line in header (skip lines containing email/phone/“resume”)
    header = lines[:10]
    header = [l for l in header if ("@" not in l and not _PHONE_RE.search(l) and "resume" not in l.lower())]
    if header:
        out["name"] = header[0][:80]

    # Company & designation heuristic:
    for l in lines:
        if "," in l and len(l) <= 160:
            left, right = [s.strip() for s in l.split(",", 1)]
            if 2 <= len(left) <= 60 and 2 <= len(right) <= 100:
                clean_company = _strip_trailing_meta(right)
                if clean_company:
                    out["designation"] = out["designation"] or left
                    out["company"] = out["company"] or clean_company
                    if out["company"] and out["designation"]:
                        break

    # Skills: bag-of-words match
    text_l = text.lower()
    skills: List[str] = []
    for kw in _SKILL_KEYWORDS:
        if kw in text_l:
            skills.append(kw)
    out["skills"] = sorted(set(skills))[:20]

    return out

# -----------------------------------------------------------------------------
# LLM extraction (optional)
# -----------------------------------------------------------------------------

_SYSTEM_PROMPT = (
    "You are a resume parser. Extract EXACT JSON with these fields: "
    "name, email, phone, company, designation, skills (array of strings). "
    "Do not include any other keys. Phone should be digits only. "
    "If a field is unknown, use an empty string or empty array."
)

def llm_extract(text: str, hints: Dict[str, object]) -> Extracted:
    """
    Ask the configured LLM to produce structured JSON. Hints are rule-based picks.
    If provider/key is missing or request fails, returns an empty Extracted().
    """
    user = (
        "Resume text follows between <TEXT> tags. Use hints when reasonable, "
        "but correct them if obviously wrong.\n\n"
        f"HINTS: {hints}\n\n"
        f"<TEXT>\n{text}\n</TEXT>"
    )
    schema = {
        "type": "object",
        "properties": {
            "name": {"type": "string"},
            "email": {"type": "string"},
            "phone": {"type": "string"},
            "company": {"type": "string"},
            "designation": {"type": "string"},
            "skills": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["name", "email", "phone", "company", "designation", "skills"],
        "additionalProperties": False,
    }
    data = generate_structured(schema, _SYSTEM_PROMPT, user)
    if not data:
        return Extracted()
    try:
        return Extracted(**data)
    except Exception:
        return Extracted()

# -----------------------------------------------------------------------------
# Merge & update helpers
# -----------------------------------------------------------------------------

def merge_results(rule: Dict[str, object], llm: Extracted | None) -> Tuple[Extracted, Dict[str, float]]:
    """
    Merge rule-based dict with optional LLM result into an Extracted + confidences.
    Rules:
      - email/phone: prefer deterministic if present (conf=1.0); else LLM (0.7).
      - name/company/designation: prefer LLM if present; else rule (0.6).
      - skills: union (cap 20).
    """
    base = Extracted(**{**Extracted().model_dump(), **(rule or {})})
    llm = llm or Extracted()

    out = Extracted(
        name = llm.name or base.name,
        email = base.email or llm.email,
        phone = base.phone or llm.phone,
        company = llm.company or base.company,
        designation = llm.designation or base.designation,
        skills = sorted(set((base.skills or []) + (llm.skills or [])))[:20],
    )

    conf = Confidence(
        name = 0.8 if llm.name else (0.6 if base.name else 0.0),
        email = 1.0 if base.email else (0.7 if llm.email else 0.0),
        phone = 1.0 if base.phone else (0.7 if llm.phone else 0.0),
        company = 0.8 if llm.company else (0.6 if base.company else 0.0),
        designation = 0.8 if llm.designation else (0.6 if base.designation else 0.0),
        skills = 0.9 if out.skills else 0.0,
    ).to_dict()

    for k in conf:
        conf[k] = clamp01(conf[k])

    return out, conf

def update_candidate(candidate, extracted: Extracted) -> None:
    """Apply extracted values to Candidate; keep existing values if they’re already set and better."""
    dirty = False

    def set_if(field: str, new_val: str):
        nonlocal dirty
        cur = getattr(candidate, field)
        if not cur and new_val:
            setattr(candidate, field, new_val)
            dirty = True

    set_if("name", extracted.name)
    set_if("email", extracted.email)
    set_if("phone", extracted.phone)
    set_if("company", extracted.company)
    set_if("designation", extracted.designation)

    if not candidate.skills and extracted.skills:
        candidate.skills = extracted.skills
        dirty = True

    if dirty:
        candidate.save(update_fields=[
            "name","email","phone","company","designation","skills","updated_at"
        ])
